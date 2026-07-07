import json
import os
import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

# --- Colores de Marca ---
HEX_PRIMARY = "572D55"  # Morado Berenjena
HEX_ACCENT = "A6225F"   # Rosa Intenso
HEX_TEXT = "333333"     # Gris Oscuro
HEX_LIGHT_BG = "F5EFF4" # Lavanda muy claro
HEX_WHITE = "FFFFFF"
HEX_BORDER = "D3C9D1"

COLOR_PRIMARY = RGBColor(87, 45, 85)
COLOR_ACCENT = RGBColor(166, 34, 95)
COLOR_TEXT = RGBColor(51, 51, 51)
COLOR_WHITE = RGBColor(255, 255, 255)

# --- Helper XML Functions for python-docx Styling ---
def set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=140, bottom=140, left=180, right=180): # in dxa (1 pt = 20 dxa)
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'''<w:tcMar {nsdecls("w")}>
        <w:top w:w="{top}" w:type="dxa"/>
        <w:bottom w:w="{bottom}" w:type="dxa"/>
        <w:left w:w="{left}" w:type="dxa"/>
        <w:right w:w="{right}" w:type="dxa"/>
    </w:tcMar>''')
    tcPr.append(tcMar)

def set_callout_borders(cell, border_color):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'''<w:tcBorders {nsdecls("w")}>
        <w:top w:val="none"/>
        <w:left w:val="single" w:sz="36" w:space="0" w:color="{border_color}"/>
        <w:bottom w:val="none"/>
        <w:right w:val="none"/>
    </w:tcBorders>''')
    tcPr.append(tcBorders)

def set_table_borders(table, color=HEX_BORDER):
    tblPr = table._tbl.tblPr
    borders = parse_xml(f'''<w:tblBorders {nsdecls("w")}>
        <w:top w:val="single" w:sz="6" w:space="0" w:color="{color}"/>
        <w:left w:val="none"/>
        <w:bottom w:val="single" w:sz="6" w:space="0" w:color="{color}"/>
        <w:right w:val="none"/>
        <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{color}"/>
        <w:insideV w:val="none"/>
    </w:tblBorders>''')
    tblPr.append(borders)

def add_formatted_text(paragraph, text, default_color=COLOR_TEXT, italic=False):
    # Procesa formato básico de negrita **
    if not text:
        return
    parts = text.split("**")
    for idx, part in enumerate(parts):
        if not part and idx == 0:
            continue
        run = paragraph.add_run(part)
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        run.font.color.rgb = default_color
        run.italic = italic
        if idx % 2 == 1:
            run.bold = True

def add_paragraph_with_spacing(doc, text="", style='Normal', space_after=6, space_before=0, bullet=False):
    style_name = 'List Bullet' if bullet else style
    p = doc.add_paragraph(style=style_name)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = 1.15
    if text:
        add_formatted_text(p, text)
    return p

def add_custom_heading(doc, text, level, space_before=12, space_after=6, keep_with_next=True):
    h = doc.add_heading(level=level)
    h.paragraph_format.space_before = Pt(space_before)
    h.paragraph_format.space_after = Pt(space_after)
    h.paragraph_format.keep_with_next = keep_with_next
    
    run = h.add_run(text)
    run.font.name = 'Calibri Light' if level == 1 else 'Calibri'
    run.bold = True
    
    if level == 1:
        run.font.size = Pt(20)
        run.font.color.rgb = COLOR_PRIMARY
    elif level == 2:
        run.font.size = Pt(15)
        run.font.color.rgb = COLOR_ACCENT
    elif level == 3:
        run.font.size = Pt(12)
        run.font.color.rgb = COLOR_PRIMARY
    return h

def add_callout(doc, text, title=""):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    # Hacer que la tabla ocupe todo el ancho
    table.columns[0].width = Inches(6.5)
    
    cell = table.cell(0, 0)
    set_cell_background(cell, HEX_LIGHT_BG)
    set_cell_margins(cell, top=160, bottom=160, left=200, right=200)
    set_callout_borders(cell, HEX_ACCENT)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    
    if title:
        run_title = p.add_run(title + "\n")
        run_title.font.name = 'Calibri'
        run_title.font.size = Pt(11)
        run_title.bold = True
        run_title.font.color.rgb = COLOR_ACCENT
        
    add_formatted_text(p, text, default_color=COLOR_PRIMARY, italic=True)
    add_paragraph_with_spacing(doc, space_after=12) # Espacio después de la tabla

# --- Main Generator ---
def main():
    # Cargar datos consolidados
    script_dir = os.path.dirname(os.path.abspath(__file__))
    json_path = os.path.join(script_dir, "../../data.json")
    
    with open(json_path, 'r', encoding='utf-8') as f:
        data = json.load(f)
        
    doctor = data["doctor"]
    services = data["services"]
    symptoms = data["symptoms"]
    diseases = data["diseases"]
    transfusion = data["transfusion"]
    privacy = data["privacyPolicy"]
    
    # ----------------------------------------------------
    # 1. GENERACIÓN DEL DOCUMENTO WORD (.docx)
    # ----------------------------------------------------
    doc = Document()
    
    # Configurar márgenes
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        # Cabecera y pie
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False
        
        # Agregar número de página en el pie de página
        footer = section.footer
        f_p = footer.paragraphs[0]
        f_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        f_run = f_p.add_run(f"Dra. Lizbeth Yamilet Hernández Verdugo  |  Contenido del Sitio Web")
        f_run.font.name = 'Calibri'
        f_run.font.size = Pt(9)
        f_run.font.color.rgb = RGBColor(120, 120, 120)

    # 1.1 Portada Elegante
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(80)
    title_p.paragraph_format.space_after = Pt(20)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    t_run = title_p.add_run("GUÍA COMPLETA DE CONTENIDO WEB")
    t_run.font.name = 'Calibri Light'
    t_run.font.size = Pt(28)
    t_run.bold = True
    t_run.font.color.rgb = COLOR_PRIMARY
    
    # Línea decorativa
    line_p = doc.add_paragraph()
    line_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    line_p.paragraph_format.space_after = Pt(20)
    l_run = line_p.add_run("―" * 25)
    l_run.font.color.rgb = COLOR_ACCENT
    l_run.bold = True
    
    doc_title_p = doc.add_paragraph()
    doc_title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc_title_p.paragraph_format.space_after = Pt(120)
    dt_run = doc_title_p.add_run("Estructura de Textos, Servicios, Síntomas, Enfermedades y Medicina Transfusional")
    dt_run.font.name = 'Calibri'
    dt_run.font.size = Pt(14)
    dt_run.font.color.rgb = COLOR_TEXT
    
    # Info de la Doctora en la portada
    info_p = doc.add_paragraph()
    info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_p.paragraph_format.space_after = Pt(6)
    dr_name_run = info_p.add_run(f"{doctor['title']} {doctor['name']}\n")
    dr_name_run.font.name = 'Calibri'
    dr_name_run.font.size = Pt(16)
    dr_name_run.bold = True
    dr_name_run.font.color.rgb = COLOR_PRIMARY
    
    dr_title_run = info_p.add_run(f"{doctor['specialty']}\n")
    dr_title_run.font.name = 'Calibri'
    dr_title_run.font.size = Pt(11)
    dr_title_run.font.color.rgb = COLOR_TEXT
    
    dr_sub_run = info_p.add_run(f"Alta Especialidad en {doctor['subspecialty']}")
    dr_sub_run.font.name = 'Calibri'
    dr_sub_run.font.size = Pt(11)
    dr_sub_run.font.color.rgb = COLOR_TEXT
    
    # Fecha de generación
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.paragraph_format.space_before = Pt(40)
    date_run = date_p.add_run(f"México, D.F.  |  {datetime.date.today().strftime('%d de %B de %Y')}")
    date_run.font.name = 'Calibri'
    date_run.font.size = Pt(10)
    date_run.font.color.rgb = RGBColor(125, 125, 125)
    
    doc.add_page_break()

    # 1.2 Sección 1: Perfil de la Doctora
    add_custom_heading(doc, "1. Perfil Profesional de la Doctora", level=1, space_before=18, space_after=12)
    add_paragraph_with_spacing(doc, doctor["bio"])
    
    add_custom_heading(doc, "Filosofía Médica", level=2, space_before=12, space_after=6)
    add_paragraph_with_spacing(doc, doctor["philosophy"])
    
    add_custom_heading(doc, "Datos del Consultorio e Información General", level=2, space_before=12, space_after=6)
    
    # Crear tabla de información general del consultorio
    info_table = doc.add_table(rows=0, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(info_table)
    
    general_details = [
      ("Dirección del Consultorio", doctor["address"]),
      ("Teléfono de Contacto", doctor["phone"]),
      ("WhatsApp de Citas", doctor["whatsapp"]),
      ("Correo Electrónico", doctor["email"]),
      ("Costo de Consulta", f"${doctor['consultationPrice']} MXN"),
      ("Métodos de Pago", ", ".join(doctor["paymentMethods"])),
      ("Horarios de Atención", doctor["schedule"]),
      ("Cédula Profesional", doctor["cedula"]),
      ("Cédula de Especialidad", doctor["cedulaEspecialidad"]),
      ("Cédula de Subespecialidad", doctor["cedulaSubespecialidad"]),
      ("Registro COFEPRIS", doctor["cofepris"]),
    ]
    
    for label, val in general_details:
        row = info_table.add_row()
        cell_lbl = row.cells[0]
        cell_val = row.cells[1]
        
        cell_lbl.width = Inches(2.2)
        cell_val.width = Inches(4.3)
        
        set_cell_background(cell_lbl, HEX_LIGHT_BG)
        set_cell_margins(cell_lbl)
        set_cell_margins(cell_val)
        
        p_lbl = cell_lbl.paragraphs[0]
        p_lbl.paragraph_format.space_after = Pt(0)
        p_lbl.paragraph_format.line_spacing = 1.15
        run_lbl = p_lbl.add_run(label)
        run_lbl.font.name = 'Calibri'
        run_lbl.bold = True
        run_lbl.font.color.rgb = COLOR_PRIMARY
        
        p_val = cell_val.paragraphs[0]
        p_val.paragraph_format.space_after = Pt(0)
        p_val.paragraph_format.line_spacing = 1.15
        run_val = p_val.add_run(val)
        run_val.font.name = 'Calibri'
        run_val.font.color.rgb = COLOR_TEXT
        
    add_paragraph_with_spacing(doc, space_after=12) # Espacio después de la tabla
    
    # Formación académica
    add_custom_heading(doc, "Educación y Formación Académica", level=2, space_before=12, space_after=6)
    for edu in doctor["education"]:
        p = add_paragraph_with_spacing(doc, bullet=True, space_after=4)
        run_deg = p.add_run(f"**{edu['degree']}** — {edu['institution']}")
        run_deg.font.name = 'Calibri'
        run_deg.font.color.rgb = COLOR_TEXT
        if edu.get("year"):
            p.add_run(f" ({edu['year']})")
            
    # Certificaciones
    add_custom_heading(doc, "Certificaciones y Colegios Médicos", level=2, space_before=12, space_after=6)
    for cert in doctor["certifications"]:
        p = add_paragraph_with_spacing(doc, bullet=True, space_after=4)
        run_cert = p.add_run(f"**{cert['name']}** — {cert['institution']}")
        run_cert.font.name = 'Calibri'
        run_cert.font.color.rgb = COLOR_TEXT
        if cert.get("year"):
            p.add_run(f" ({cert['year']})")
            
    # Experiencia
    add_custom_heading(doc, "Trayectoria y Experiencia", level=2, space_before=12, space_after=6)
    for exp in doctor["experience"]:
        p = add_paragraph_with_spacing(doc, space_after=6)
        run_exp = p.add_run(f"• **{exp['year']} - {exp['title']}:** ")
        run_exp.font.name = 'Calibri'
        run_exp.font.color.rgb = COLOR_TEXT
        p.add_run(exp["description"])

    doc.add_page_break()

    # 1.3 Sección 2: Servicios Médicos
    add_custom_heading(doc, "2. Portafolio de Servicios y Procedimientos Médicos", level=1, space_before=18, space_after=12)
    
    for idx, s in enumerate(services):
        add_custom_heading(doc, f"Servicio {idx+1}: {s['name']}", level=2, space_before=14, space_after=8)
        
        # Resumen y descripción larga
        add_paragraph_with_spacing(doc, f"**Descripción corta:** {s['description']}")
        add_paragraph_with_spacing(doc, f"**Descripción detallada:** {s['longDescription']}")
        
        if s.get("painDescription"):
            add_paragraph_with_spacing(doc, f"**Experiencia del Dolor:** {s['painDescription']}")
            
        # Ficha Técnica en Tabla
        add_custom_heading(doc, "Ficha Técnica del Procedimiento", level=3, space_before=10, space_after=4)
        spec_table = doc.add_table(rows=0, cols=2)
        spec_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(spec_table)
        
        specs = [
          ("Tipo de Servicio", s["type"].upper()),
          ("Tipo de Anestesia", s["anesthesiaType"]),
          ("Duración Estimada", s["duration"]),
          ("Tiempo de Recuperación", s["recoveryTime"]),
          ("¿Es doloroso?", "Sí" if s["isPainful"] else "No"),
          ("Rango de Precios", s["priceRange"]),
        ]
        
        # Agregar campos de especificaciones técnicas personalizadas
        for k, v in s["technicalSpecs"].items():
            specs.append((k, v))
            
        for label, val in specs:
            row = spec_table.add_row()
            cell_lbl = row.cells[0]
            cell_val = row.cells[1]
            
            cell_lbl.width = Inches(2.2)
            cell_val.width = Inches(4.3)
            
            set_cell_background(cell_lbl, HEX_LIGHT_BG)
            set_cell_margins(cell_lbl)
            set_cell_margins(cell_val)
            
            p_lbl = cell_lbl.paragraphs[0]
            p_lbl.paragraph_format.space_after = Pt(0)
            p_lbl.paragraph_format.line_spacing = 1.15
            run_lbl = p_lbl.add_run(label)
            run_lbl.font.name = 'Calibri'
            run_lbl.bold = True
            run_lbl.font.color.rgb = COLOR_PRIMARY
            
            p_val = cell_val.paragraphs[0]
            p_val.paragraph_format.space_after = Pt(0)
            p_val.paragraph_format.line_spacing = 1.15
            run_val = p_val.add_run(val)
            run_val.font.name = 'Calibri'
            run_val.font.color.rgb = COLOR_TEXT
            
        add_paragraph_with_spacing(doc, space_after=12)
        
        # Indicaciones y Beneficios
        add_custom_heading(doc, s["benefitsTitle"], level=3, space_before=10, space_after=4)
        if s.get("benefitsIntro"):
            add_paragraph_with_spacing(doc, s["benefitsIntro"])
        for b in s["benefits"]:
            add_paragraph_with_spacing(doc, b, bullet=True, space_after=4)
            
        # Recomendaciones posoperatorias
        add_custom_heading(doc, s["recommendationsTitle"], level=3, space_before=10, space_after=4)
        for rec in s["postOpRecommendations"]:
            add_paragraph_with_spacing(doc, rec, bullet=True, space_after=4)
            
        # Riesgos
        if s.get("risks"):
            add_custom_heading(doc, s["risksTitle"], level=3, space_before=10, space_after=4)
            for risk in s["risks"]:
                add_paragraph_with_spacing(doc, risk, bullet=True, space_after=4)
                
        # Preguntas adicionales de este servicio
        if s.get("additionalQuestions"):
            add_custom_heading(doc, "Preguntas Frecuentes del Servicio", level=3, space_before=10, space_after=4)
            for q_item in s["additionalQuestions"]:
                add_paragraph_with_spacing(doc, f"¿{q_item['question']}?", style='Normal', space_before=4, space_after=2)
                # Poner la respuesta en cursiva y con indentación
                p_ans = add_paragraph_with_spacing(doc, space_after=6)
                p_ans.paragraph_format.left_indent = Inches(0.25)
                run_ans = p_ans.add_run(q_item["answer"])
                run_ans.font.name = 'Calibri'
                run_ans.font.size = Pt(10)
                run_ans.font.color.rgb = COLOR_TEXT
                run_ans.italic = True

        # CTA Callout
        if s.get("ctaQuestion"):
            add_callout(doc, s["ctaAnswer"], title=s["ctaQuestion"])
            
        doc.add_page_break()

    # 1.4 Sección 3: Síntomas Comunes en Hematología Pediátrica
    add_custom_heading(doc, "3. Guía de Síntomas y Manifestaciones Clínicas", level=1, space_before=18, space_after=12)
    add_paragraph_with_spacing(doc, "Esta sección describe los principales motivos de consulta y síntomas sanguíneos en niños, sus posibles causas asociadas y los criterios clínicos para buscar atención especializada inmediata.")
    
    for idx, sym in enumerate(symptoms):
        add_custom_heading(doc, f"Síntoma {idx+1}: {sym['name']}", level=2, space_before=14, space_after=8)
        add_paragraph_with_spacing(doc, sym["description"])
        
        # Manifestaciones
        add_custom_heading(doc, sym["manifestacionesTitle"], level=3, space_before=10, space_after=4)
        for man in sym["manifestaciones"]:
            add_paragraph_with_spacing(doc, man, bullet=True, space_after=4)
            
        # Causas
        add_custom_heading(doc, sym["causesTitle"], level=3, space_before=10, space_after=4)
        if sym.get("causesIntro"):
            add_paragraph_with_spacing(doc, sym["causesIntro"])
        for cause in sym["causes"]:
            add_paragraph_with_spacing(doc, cause, bullet=True, space_after=4)
            
        # Signos de alarma / Cuándo consultar
        add_custom_heading(doc, sym["signosAlarmaTitle"] if sym.get("signosAlarmaTitle") else "¿Cuándo consultar al especialista?", level=3, space_before=10, space_after=4)
        add_paragraph_with_spacing(doc, sym["whyConsult"])
        
        if sym.get("signosAlarma"):
            for alarm in sym["signosAlarma"]:
                p_alarm = add_paragraph_with_spacing(doc, bullet=True, space_after=4)
                run_alarm = p_alarm.add_run(alarm)
                run_alarm.font.name = 'Calibri'
                run_alarm.font.color.rgb = COLOR_ACCENT
                run_alarm.bold = True
                
        # Pregunta principal / CTA
        if sym.get("preguntaPrincipal"):
            add_callout(doc, sym.get("preguntaPrincipalResponse", "Agende una valoración para descartar problemas mayores."), title=sym["preguntaPrincipal"])
            
        doc.add_page_break()

    # 1.5 Sección 4: Enfermedades Hematológicas
    add_custom_heading(doc, "4. Enfermedades y Condiciones Hematológicas", level=1, space_before=18, space_after=12)
    add_paragraph_with_spacing(doc, "Guía clínica sobre las principales patologías de la sangre y el sistema inmunológico en pacientes pediátricos.")
    
    for idx, d in enumerate(diseases):
        add_custom_heading(doc, f"Condición {idx+1}: {d['name']}", level=2, space_before=14, space_after=8)
        if d.get("technicalName"):
            p_tech = add_paragraph_with_spacing(doc, space_after=6)
            run_tech = p_tech.add_run(f"Nombre Clínico/Técnico: {d['technicalName']}")
            run_tech.font.name = 'Calibri'
            run_tech.bold = True
            run_tech.font.color.rgb = COLOR_ACCENT
            
        add_paragraph_with_spacing(doc, d["description"])
        
        # Categoría
        p_cat = add_paragraph_with_spacing(doc, space_after=10)
        run_cat_lbl = p_cat.add_run("Clasificación Médica: ")
        run_cat_lbl.font.color.rgb = COLOR_PRIMARY
        run_cat_lbl.bold = True
        p_cat.add_run(d["category"].replace("-", " ").title())
        
        # Síntomas de la enfermedad
        add_custom_heading(doc, "Síntomas Comunes en Niños", level=3, space_before=10, space_after=4)
        for sym_disease in d["symptoms"]:
            add_paragraph_with_spacing(doc, sym_disease, bullet=True, space_after=4)
            
        # Causas
        add_custom_heading(doc, "Causas y Origen", level=3, space_before=10, space_after=4)
        for cause in d["causes"]:
            add_paragraph_with_spacing(doc, cause, bullet=True, space_after=4)
            
        # Factores de riesgo
        if d.get("riskFactors"):
            add_custom_heading(doc, "Factores de Riesgo", level=3, space_before=10, space_after=4)
            for risk in d["riskFactors"]:
                add_paragraph_with_spacing(doc, risk, bullet=True, space_after=4)
                
        # Estadísticas en México
        if d.get("mexicoStats"):
            add_custom_heading(doc, "Estadísticas y Panorama en México", level=3, space_before=10, space_after=4)
            add_paragraph_with_spacing(doc, d["mexicoStats"])
            
        # Complicaciones
        if d.get("complications"):
            add_custom_heading(doc, "Complicaciones Potenciales", level=3, space_before=10, space_after=4)
            for comp in d["complications"]:
                add_paragraph_with_spacing(doc, comp, bullet=True, space_after=4)
                
        # Tratamientos
        add_custom_heading(doc, "Abordaje Terapéutico (Tratamientos)", level=3, space_before=10, space_after=4)
        for treat in d["treatments"]:
            add_paragraph_with_spacing(doc, treat, bullet=True, space_after=4)
            
        # Diagnósticos diferenciales o descartes
        if d.get("diseasesToRuleOut"):
            add_custom_heading(doc, "Diagnósticos Diferenciales a Descartar", level=3, space_before=10, space_after=4)
            for rule_out in d["diseasesToRuleOut"]:
                add_paragraph_with_spacing(doc, rule_out, bullet=True, space_after=4)
                
        # Cuándo consultar
        if d.get("whenToConsult"):
            add_custom_heading(doc, "¿Cuándo buscar una valoración con Hematología Pediátrica?", level=3, space_before=10, space_after=4)
            for w in d["whenToConsult"]:
                add_paragraph_with_spacing(doc, w, bullet=True, space_after=4)
                
        # FAQs
        add_custom_heading(doc, "Preguntas Frecuentes (FAQs)", level=3, space_before=10, space_after=4)
        for faq in d["faqs"]:
            add_paragraph_with_spacing(doc, f"¿{faq['question']}?", style='Normal', space_before=6, space_after=2)
            p_ans = add_paragraph_with_spacing(doc, space_after=6)
            p_ans.paragraph_format.left_indent = Inches(0.25)
            run_ans = p_ans.add_run(faq["answer"])
            run_ans.font.name = 'Calibri'
            run_ans.font.size = Pt(10)
            run_ans.font.color.rgb = COLOR_TEXT
            run_ans.italic = True
            
        # Fuentes científicas
        if d.get("sources"):
            add_custom_heading(doc, "Fuentes de Referencia Médica", level=3, space_before=10, space_after=4)
            for source in d["sources"]:
                p_s = add_paragraph_with_spacing(doc, bullet=True, space_after=2)
                run_s = p_s.add_run(source)
                run_s.font.size = Pt(9.5)
                run_s.font.color.rgb = RGBColor(120, 120, 120)
                
        doc.add_page_break()

    # 1.6 Sección 5: Medicina Transfusional
    add_custom_heading(doc, "5. Medicina Transfusional y Aféresis Terapéutica", level=1, space_before=18, space_after=12)
    add_paragraph_with_spacing(doc, "La Medicina Transfusional Pediátrica y las terapias de aféresis constituyen un pilar fundamental en el soporte y curación de enfermedades hematológicas complejas, oncológicas e inmunológicas. Esta sección compila las guías del servicio.")
    
    for t in transfusion:
        add_custom_heading(doc, t["title"], level=2, space_before=14, space_after=8)
        
        if t.get("intro"):
            add_paragraph_with_spacing(doc, t["intro"])
            
        # Preguntas y respuestas principales
        add_custom_heading(doc, "Guía de Información y Criterios", level=3, space_before=10, space_after=4)
        for sec in t["sections"]:
            add_paragraph_with_spacing(doc, sec["q"], style='Normal', space_before=6, space_after=2)
            p_ans = add_paragraph_with_spacing(doc, space_after=6)
            p_ans.paragraph_format.left_indent = Inches(0.25)
            add_formatted_text(p_ans, sec["a"])
            p_ans.runs[0].font.size = Pt(10)
            p_ans.runs[0].italic = True
            
        # Posibles causas o indicaciones de aféresis/hierro
        if t.get("posiblesCausas"):
            add_custom_heading(doc, "Criterios e Indicaciones Médicas Asociadas", level=3, space_before=10, space_after=4)
            for pc in t["posiblesCausas"]:
                add_paragraph_with_spacing(doc, pc, bullet=True, space_after=4)
                
        # Pregunta de Alerta / Callout
        if t.get("alertQuestion"):
            add_callout(doc, t["alertQuestion"]["a"], title=t["alertQuestion"]["q"])
            
        # Fuentes de medicina transfusional
        if t.get("sources"):
            add_custom_heading(doc, "Fuentes Bibliográficas", level=3, space_before=10, space_after=4)
            for src in t["sources"]:
                p_s = add_paragraph_with_spacing(doc, bullet=True, space_after=2)
                run_s = p_s.add_run(src)
                run_s.font.size = Pt(9.5)
                run_s.font.color.rgb = RGBColor(120, 120, 120)
                
        doc.add_page_break()

    # 1.7 Sección 6: Páginas Adicionales
    add_custom_heading(doc, "6. Páginas de Soporte Legal y de Contacto", level=1, space_before=18, space_after=12)
    
    # Aviso de privacidad
    add_custom_heading(doc, privacy["title"], level=2, space_before=12, space_after=6)
    add_paragraph_with_spacing(doc, privacy["intro"])
    
    for priv_sec in privacy["sections"]:
        add_custom_heading(doc, priv_sec["title"], level=3, space_before=10, space_after=4)
        add_paragraph_with_spacing(doc, priv_sec["content"])
        
    p_upd = add_paragraph_with_spacing(doc, space_before=12)
    run_upd = p_upd.add_run(f"Última actualización: {privacy['lastUpdate']}")
    run_upd.font.size = Pt(9)
    run_upd.italic = True
    run_upd.font.color.rgb = RGBColor(120, 120, 120)
    
    doc.add_page_break()
    
    # Página de Contacto
    add_custom_heading(doc, "Ubicación del Consultorio y Datos de Contacto", level=2, space_before=12, space_after=6)
    add_paragraph_with_spacing(doc, "Resumen del contenido textual de la página de contacto:")
    
    contact_table = doc.add_table(rows=0, cols=2)
    contact_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(contact_table)
    
    contact_details = [
      ("Doctora", f"{doctor['title']} {doctor['name']}"),
      ("Dirección", f"{doctor['address']}, {doctor['city']}, {doctor['state']}"),
      ("WhatsApp Citas", doctor["whatsapp"]),
      ("Teléfono Emergencias", doctor["phone"]),
      ("Email Citas", doctor["email"]),
      ("Horarios de Consultorio", doctor["schedule"]),
      ("Inversión de Consulta", f"${doctor['consultationPrice']} MXN"),
      ("Métodos de Pago Aceptados", ", ".join(doctor["paymentMethods"])),
      ("Enlace de Google Maps", doctor["googleMapsUrl"])
    ]
    
    for label, val in contact_details:
        row = contact_table.add_row()
        cell_lbl = row.cells[0]
        cell_val = row.cells[1]
        
        cell_lbl.width = Inches(2.2)
        cell_val.width = Inches(4.3)
        
        set_cell_background(cell_lbl, HEX_LIGHT_BG)
        set_cell_margins(cell_lbl)
        set_cell_margins(cell_val)
        
        p_lbl = cell_lbl.paragraphs[0]
        p_lbl.paragraph_format.space_after = Pt(0)
        p_lbl.paragraph_format.line_spacing = 1.15
        run_lbl = p_lbl.add_run(label)
        run_lbl.font.name = 'Calibri'
        run_lbl.bold = True
        run_lbl.font.color.rgb = COLOR_PRIMARY
        
        p_val = cell_val.paragraphs[0]
        p_val.paragraph_format.space_after = Pt(0)
        p_val.paragraph_format.line_spacing = 1.15
        run_val = p_val.add_run(val)
        run_val.font.name = 'Calibri'
        run_val.font.color.rgb = COLOR_TEXT
        
    add_paragraph_with_spacing(doc, space_after=12)
    
    # Guardar Word Document
    docx_output = os.path.join(script_dir, "../../Dra_Lizbeth_Verdugo_Contenido.docx")
    doc.save(docx_output)
    print(f"Documento Word (.docx) guardado correctamente en: {docx_output}")

    # ----------------------------------------------------
    # 2. GENERACIÓN DEL DOCUMENTO HTML (.html)
    # ----------------------------------------------------
    # Genera un HTML con diseño responsivo premium, fuentes de Google e interacciones fluidas.
    html_output = os.path.join(script_dir, "../../Dra_Lizbeth_Verdugo_Contenido.html")
    
    html_content = f"""<!DOCTYPE html>
<html lang="es">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Dra. Lizbeth Hernández | Contenido Web</title>
    <link href="https://fonts.googleapis.com/css2?family=Outfit:wght@300;400;600;700;800&family=Playfair+Display:ital,wght@0,600;0,800;1,500&display=swap" rel="stylesheet">
    <style>
        :root {{
            --primary: #{HEX_PRIMARY};
            --accent: #{HEX_ACCENT};
            --text: #{HEX_TEXT};
            --bg: #FDFBFC;
            --card-bg: #FFFFFF;
            --light-bg: #{HEX_LIGHT_BG};
            --border: #{HEX_BORDER};
            --primary-light: rgba(87, 45, 85, 0.05);
            --accent-light: rgba(166, 34, 95, 0.04);
        }}
        * {{
            box-sizing: border-box;
            margin: 0;
            padding: 0;
        }}
        body {{
            font-family: 'Outfit', sans-serif;
            background-color: var(--bg);
            color: var(--text);
            line-height: 1.6;
            display: flex;
            min-height: 100vh;
        }}
        
        /* Sidebar de Navegación */
        aside {{
            width: 320px;
            background-color: #FFFFFF;
            border-right: 1px solid var(--border);
            padding: 2rem;
            position: fixed;
            top: 0;
            bottom: 0;
            left: 0;
            overflow-y: auto;
            display: flex;
            flex-col: column;
            flex-direction: column;
            gap: 1.5rem;
            z-index: 100;
        }}
        .aside-header h1 {{
            font-family: 'Playfair Display', serif;
            font-size: 1.4rem;
            font-weight: 800;
            color: var(--primary);
            margin-bottom: 0.2rem;
        }}
        .aside-header p {{
            font-size: 0.8rem;
            font-weight: 600;
            text-transform: uppercase;
            letter-spacing: 1px;
            color: var(--accent);
        }}
        .nav-links {{
            display: flex;
            flex-direction: column;
            gap: 0.5rem;
            list-style: none;
            margin-top: 1rem;
        }}
        .nav-links a {{
            display: block;
            padding: 0.75rem 1rem;
            border-radius: 1rem;
            color: var(--text);
            text-decoration: none;
            font-weight: 600;
            font-size: 0.9rem;
            transition: all 0.2s ease;
            border: 1px solid transparent;
        }}
        .nav-links a:hover {{
            background-color: var(--primary-light);
            color: var(--primary);
        }}
        .nav-links .active {{
            background-color: var(--primary);
            color: #FFFFFF;
        }}

        /* Contenido Principal */
        main {{
            margin-left: 320px;
            flex: 1;
            padding: 3rem 4rem;
            max-width: 1000px;
        }}
        header.doc-header {{
            margin-bottom: 3rem;
            border-bottom: 2px solid var(--primary-light);
            padding-bottom: 2rem;
        }}
        .doc-title {{
            font-family: 'Playfair Display', serif;
            font-size: 3rem;
            font-weight: 800;
            color: var(--primary);
            line-height: 1.1;
            margin-bottom: 1rem;
        }}
        .doc-subtitle {{
            font-size: 1.1rem;
            color: var(--accent);
            font-weight: 600;
            text-transform: uppercase;
            letter-spacing: 1.5px;
        }}
        
        section {{
            margin-bottom: 4rem;
            scroll-margin-top: 2rem;
        }}
        h2 {{
            font-family: 'Playfair Display', serif;
            font-size: 2rem;
            color: var(--primary);
            margin-bottom: 1.5rem;
            border-bottom: 2px solid var(--accent);
            padding-bottom: 0.5rem;
            display: inline-block;
        }}
        h3 {{
            font-size: 1.3rem;
            color: var(--accent);
            margin-top: 2rem;
            margin-bottom: 1rem;
            font-weight: 700;
            text-transform: uppercase;
            letter-spacing: 0.5px;
        }}
        h4 {{
            font-size: 1.1rem;
            color: var(--primary);
            margin-top: 1.5rem;
            margin-bottom: 0.5rem;
            font-weight: 700;
        }}
        p {{
            margin-bottom: 1.2rem;
            font-size: 1rem;
            color: #4a4a4a;
            text-align: justify;
        }}
        ul, ol {{
            margin-bottom: 1.5rem;
            padding-left: 1.5rem;
        }}
        li {{
            margin-bottom: 0.5rem;
            color: #4a4a4a;
        }}
        
        /* Tablas */
        table {{
            width: 100%;
            border-collapse: collapse;
            margin-top: 1.5rem;
            margin-bottom: 2rem;
            background-color: var(--card-bg);
            border-radius: 1.5rem;
            overflow: hidden;
            box-shadow: 0 4px 15px rgba(0,0,0,0.03);
            border: 1px solid var(--border);
        }}
        th, td {{
            padding: 1rem 1.25rem;
            text-align: left;
            border-bottom: 1px solid var(--primary-light);
            font-size: 0.95rem;
        }}
        th {{
            background-color: var(--primary);
            color: #FFFFFF;
            font-weight: 700;
            text-transform: uppercase;
            font-size: 0.8rem;
            letter-spacing: 1px;
        }}
        tr:last-child td {{
            border-bottom: none;
        }}
        tr:nth-child(even) td:first-child {{
            background-color: var(--primary-light);
        }}
        td:first-child {{
            font-weight: 700;
            color: var(--primary);
            width: 30%;
        }}
        
        /* Callouts / Tarjetas de Alerta */
        .callout {{
            background-color: var(--light-bg);
            border-left: 5px solid var(--accent);
            padding: 1.5rem 2rem;
            border-radius: 0.5rem 1.5rem 1.5rem 0.5rem;
            margin: 2rem 0;
            box-shadow: 0 5px 15px rgba(166, 34, 95, 0.03);
        }}
        .callout-title {{
            font-weight: 700;
            color: var(--accent);
            margin-bottom: 0.5rem;
            font-size: 1.05rem;
            text-transform: uppercase;
        }}
        .callout-text {{
            font-style: italic;
            color: var(--primary);
            font-size: 0.95rem;
            margin-bottom: 0;
        }}
        
        .badge {{
            display: inline-block;
            padding: 0.25rem 0.75rem;
            background-color: var(--accent-light);
            color: var(--accent);
            border-radius: 2rem;
            font-size: 0.75rem;
            font-weight: 700;
            text-transform: uppercase;
            margin-bottom: 0.5rem;
            border: 1px solid rgba(166,34,95,0.1);
        }}
        
        .faq-box {{
            background: #FFFFFF;
            border: 1px solid var(--border);
            border-radius: 1.5rem;
            padding: 1.5rem;
            margin-bottom: 1.5rem;
            box-shadow: 0 4px 10px rgba(0,0,0,0.02);
        }}
        .faq-q {{
            font-weight: 700;
            color: var(--primary);
            margin-bottom: 0.5rem;
            font-size: 1.05rem;
        }}
        .faq-a {{
            color: #555555;
            font-style: italic;
            margin-bottom: 0;
            padding-left: 0.5rem;
        }}
        
        .sources-list {{
            background: rgba(0,0,0,0.01);
            padding: 1rem 1.5rem;
            border-radius: 1rem;
            border: 1px dashed var(--border);
        }}
        .sources-list li {{
            font-size: 0.85rem;
            color: #777777;
            list-style-type: square;
        }}
        
        /* Botón de impresión o exportar */
        .btn-actions {{
            display: flex;
            gap: 1rem;
            margin-top: 1.5rem;
        }}
        .btn-print {{
            padding: 0.75rem 1.5rem;
            background-color: var(--primary);
            color: #FFFFFF;
            border: none;
            border-radius: 2rem;
            font-weight: 700;
            cursor: pointer;
            transition: all 0.3s ease;
            font-size: 0.85rem;
            box-shadow: 0 4px 10px rgba(87, 45, 85, 0.2);
        }}
        .btn-print:hover {{
            background-color: var(--accent);
            box-shadow: 0 4px 10px rgba(166, 34, 95, 0.2);
        }}
        
        @media (max-width: 1024px) {{
            body {{
                flex-direction: column;
            }}
            aside {{
                position: relative;
                width: 100%;
                border-right: none;
                border-bottom: 1px solid var(--border);
                padding: 1.5rem;
            }}
            main {{
                margin-left: 0;
                padding: 2rem;
            }}
        }}
    </style>
</head>
<body>
    <aside>
        <div class="aside-header">
            <h1>Dra. Lizbeth Hernández</h1>
            <p>Contenido del Sitio Web</p>
        </div>
        <ul class="nav-links">
            <li><a href="#perfil" class="active">1. Perfil Profesional</a></li>
            <li><a href="#servicios">2. Servicios y Procedimientos</a></li>
            <li><a href="#sintomas">3. Guía de Síntomas</a></li>
            <li><a href="#enfermedades">4. Enfermedades</a></li>
            <li><a href="#transfusion">5. Medicina Transfusional</a></li>
            <li><a href="#legales">6. Páginas de Soporte</a></li>
        </ul>
        <div class="btn-actions">
            <button class="btn-print" onclick="window.print()">Imprimir / Guardar PDF</button>
        </div>
    </aside>

    <main>
        <header class="doc-header">
            <div class="doc-subtitle">Manual de Contenidos Consolidados</div>
            <div class="doc-title">Estructura y Textos del Sitio Web</div>
            <p>Este documento interactivo consolida el total de textos, descripciones médicas, especificaciones de procedimientos e información del perfil clínico del portal de la Dra. Lizbeth Hernández Verdugo. Está diseñado para ser leído en pantalla, impreso o copiado directamente en procesadores de texto.</p>
        </header>

        <!-- SECCIÓN 1: PERFIL DOCTORA -->
        <section id="perfil">
            <h2>1. Perfil Profesional de la Doctora</h2>
            <p style="font-size: 1.1rem; font-weight: 600; color: var(--primary); margin-bottom: 1.5rem;">{doctor['title']} {doctor['name']}</p>
            <p>{doctor['bio']}</p>
            
            <h3>Filosofía Médica</h3>
            <div class="callout">
                <p class="callout-text">"{doctor['philosophy']}"</p>
            </div>
            
            <h3>Información General del Consultorio</h3>
            <table>
                <thead>
                    <tr>
                        <th>Concepto / Campo</th>
                        <th>Detalle de Información</th>
                    </tr>
                </thead>
                <tbody>
                    <tr><td>Especialidad</td><td>{doctor['specialty']}</td></tr>
                    <tr><td>Subespecialidad</td><td>Alta Especialidad en {doctor['subspecialty']}</td></tr>
                    <tr><td>Dirección del Consultorio</td><td>{doctor['address']}</td></tr>
                    <tr><td>Teléfono de Contacto</td><td>{doctor['phone']}</td></tr>
                    <tr><td>WhatsApp de Citas</td><td>{doctor['whatsapp']}</td></tr>
                    <tr><td>Correo Electrónico</td><td>{doctor['email']}</td></tr>
                    <tr><td>Costo de Consulta</td><td>${doctor['consultationPrice']} MXN</td></tr>
                    <tr><td>Métodos de Pago</td><td>{", ".join(doctor['paymentMethods'])}</td></tr>
                    <tr><td>Horarios de Atención</td><td>{doctor['schedule']}</td></tr>
                    <tr><td>Cédula Profesional</td><td>{doctor['cedula']}</td></tr>
                    <tr><td>Cédula de Especialidad</td><td>{doctor['cedulaEspecialidad']}</td></tr>
                    <tr><td>Cédula de Subespecialidad</td><td>{doctor['cedulaSubespecialidad']}</td></tr>
                    <tr><td>Registro COFEPRIS</td><td>{doctor['cofepris']}</td></tr>
                </tbody>
            </table>

            <h3>Educación y Formación</h3>
            <ul>
    """
    
    for edu in doctor["education"]:
        year_str = f" ({edu['year']})" if edu.get("year") else ""
        html_content += f"<li><strong>{edu['degree']}</strong> — {edu['institution']}{year_str}</li>"
        
    html_content += """
            </ul>

            <h3>Certificaciones</h3>
            <ul>
    """
    
    for cert in doctor["certifications"]:
        year_str = f" ({cert['year']})" if cert.get("year") else ""
        html_content += f"<li><strong>{cert['name']}</strong> — {cert['institution']}{year_str}</li>"
        
    html_content += """
            </ul>

            <h3>Experiencia Laboral y Trayectoria</h3>
            <ul>
    """
    for exp in doctor["experience"]:
        html_content += f"<li><strong>{exp['year']} - {exp['title']}:</strong> {exp['description']}</li>"
        
    html_content += """
            </ul>
        </section>

        <!-- SECCIÓN 2: SERVICIOS -->
        <section id="servicios">
            <h2>2. Portafolio de Servicios y Procedimientos Médicos</h2>
            <p>Guía de los procedimientos diagnósticos y terapéuticos especializados realizados en consultorio y ambiente hospitalario.</p>
    """
    
    for s in services:
        is_painful_str = "Sí" if s["isPainful"] else "No"
        html_content += f"""
            <div style="margin-top: 3rem; border-top: 1px solid var(--border); padding-top: 2rem;">
                <span class="badge">{s['type'].upper()}</span>
                <h3 style="margin-top: 0.5rem; font-family: 'Playfair Display', serif; text-transform: none; letter-spacing: 0; font-size: 1.6rem; color: var(--primary);">{s['name']}</h3>
                <p><strong>Resumen Diagnóstico:</strong> {s['description']}</p>
                <p><strong>Descripción Ampliada:</strong> {s['longDescription']}</p>
        """
        
        if s.get("painDescription"):
            html_content += f"<p><strong>Experiencia y Control del Dolor:</strong> {s['painDescription']}</p>"
            
        html_content += f"""
                <h4>Ficha Técnica</h4>
                <table>
                    <thead>
                        <tr>
                            <th>Especificación</th>
                            <th>Detalle Técnico</th>
                        </tr>
                    </thead>
                    <tbody>
                        <tr><td>Tipo de Procedimiento</td><td>{s['type'].upper()}</td></tr>
                        <tr><td>Tipo de Anestesia</td><td>{s['anesthesiaType']}</td></tr>
                        <tr><td>Duración Promedio</td><td>{s['duration']}</td></tr>
                        <tr><td>Tiempo de Recuperación</td><td>{s['recoveryTime']}</td></tr>
                        <tr><td>¿Provoca dolor o molestia?</td><td>{is_painful_str}</td></tr>
                        <tr><td>Costo Aproximado</td><td>{s['priceRange']}</td></tr>
        """
        
        for k, v in s["technicalSpecs"].items():
            html_content += f"<tr><td>{k}</td><td>{v}</td></tr>"
            
        html_content += f"""
                    </tbody>
                </table>

                <h4>{s['benefitsTitle']}</h4>
                <ul>
        """
        
        for b in s["benefits"]:
            html_content += f"<li>{b}</li>"
            
        html_content += f"""
                </ul>

                <h4>{s['recommendationsTitle']}</h4>
                <ul>
        """
        
        for rec in s["postOpRecommendations"]:
            html_content += f"<li>{rec}</li>"
            
        html_content += """
                </ul>
        """
        
        if s.get("risks"):
            html_content += f"""
                <h4>{s['risksTitle']}</h4>
                <ul>
            """
            for risk in s["risks"]:
                html_content += f"<li>{risk}</li>"
            html_content += "</ul>"
            
        if s.get("additionalQuestions"):
            html_content += "<h4>Preguntas Frecuentes del Procedimiento</h4>"
            for q_item in s["additionalQuestions"]:
                html_content += f"""
                <div class="faq-box">
                    <div class="faq-q">¿{q_item['question']}?</div>
                    <div class="faq-a">{q_item['answer']}</div>
                </div>
                """
                
        if s.get("ctaQuestion"):
            html_content += f"""
                <div class="callout">
                    <div class="callout-title">{s['ctaQuestion']}</div>
                    <div class="callout-text">{s['ctaAnswer']}</div>
                </div>
            """
            
        html_content += "</div>"
        
    html_content += """
        </section>

        <!-- SECCIÓN 3: SÍNTOMAS -->
        <section id="sintomas">
            <h2>3. Guía de Síntomas y Motivos de Consulta</h2>
            <p>Sección para que padres de familia y pacientes identifiquen signos de alarma relacionados con la salud de la sangre.</p>
    """
    
    for sym in symptoms:
        html_content += f"""
            <div style="margin-top: 3rem; border-top: 1px solid var(--border); padding-top: 2rem;">
                <h3 style="font-family: 'Playfair Display', serif; text-transform: none; letter-spacing: 0; font-size: 1.5rem; color: var(--primary);">{sym['name']}</h3>
                <p>{sym['description']}</p>
                
                <h4>{sym['manifestacionesTitle']}</h4>
                <ul>
        """
        for man in sym["manifestaciones"]:
            html_content += f"<li>{man}</li>"
            
        html_content += f"""
                </ul>

                <h4>{sym['causesTitle']}</h4>
        """
        if sym.get("causesIntro"):
            html_content += f"<p>{sym['causesIntro']}</p>"
        html_content += "<ul>"
        for c_sym in sym["causes"]:
            html_content += f"<li>{c_sym}</li>"
        html_content += f"""
                </ul>

                <h4>¿Cuándo se debe consultar a la especialista?</h4>
                <p>{sym['whyConsult']}</p>
        """
        
        if sym.get("signosAlarma"):
            html_content += f"""
                <div class="callout" style="background-color: rgba(166, 34, 95, 0.03); border-left-color: var(--accent);">
                    <div class="callout-title" style="color: var(--accent);">{sym['signosAlarmaTitle'] if sym.get('signosAlarmaTitle') else 'Signos de Alarma Críticos'}</div>
                    <ul>
            """
            for alarm in sym["signosAlarma"]:
                html_content += f"<li style='color: var(--accent); font-weight: 700;'>{alarm}</li>"
            html_content += "</ul></div>"
            
        if sym.get("preguntaPrincipal"):
            html_content += f"""
                <div class="callout">
                    <div class="callout-title">{sym['preguntaPrincipal']}</div>
                    <div class="callout-text">{sym.get('preguntaPrincipalResponse', 'Agende una valoración médica inmediata.')}</div>
                </div>
            """
        html_content += "</div>"
        
    html_content += """
        </section>

        <!-- SECCIÓN 4: ENFERMEDADES -->
        <section id="enfermedades">
            <h2>4. Enfermedades y Condiciones de la Sangre</h2>
            <p>Información detallada sobre las condiciones hematológicas y del sistema inmune en niños y adolescentes.</p>
    """
    
    for d in diseases:
        tech_title = f"<span class='badge' style='background-color: var(--primary-light); color: var(--primary);'>{d['technicalName']}</span>" if d.get("technicalName") else ""
        html_content += f"""
            <div style="margin-top: 3rem; border-top: 1px solid var(--border); padding-top: 2rem;">
                <span class="badge" style="background-color: var(--primary-light); color: var(--primary);">{d['category'].replace('-', ' ').title()}</span>
                <h3 style="margin-top: 0.5rem; font-family: 'Playfair Display', serif; text-transform: none; letter-spacing: 0; font-size: 1.6rem; color: var(--primary);">{d['name']}</h3>
                {tech_title}
                <p style="margin-top: 1rem;">{d['description']}</p>
                
                <h4>Síntomas Comunes</h4>
                <ul>
        """
        for sym_d in d["symptoms"]:
            html_content += f"<li>{sym_d}</li>"
            
        html_content += """
                </ul>

                <h4>Causas y Origen</h4>
                <ul>
        """
        for cause in d["causes"]:
            html_content += f"<li>{cause}</li>"
            
        html_content += "</ul>"
        
        if d.get("riskFactors"):
            html_content += """
                <h4>Factores de Riesgo</h4>
                <ul>
            """
            for risk in d["riskFactors"]:
                html_content += f"<li>{risk}</li>"
            html_content += "</ul>"
            
        if d.get("mexicoStats"):
            html_content += f"""
                <h4>Panorama Clínico en México</h4>
                <p>{d['mexicoStats']}</p>
            """
            
        if d.get("complications"):
            html_content += """
                <h4>Complicaciones Potenciales</h4>
                <ul>
            """
            for comp in d["complications"]:
                html_content += f"<li>{comp}</li>"
            html_content += "</ul>"
            
        html_content += """
                <h4>Tratamiento y Abordaje</h4>
                <ul>
        """
        for treat in d["treatments"]:
            html_content += f"<li>{treat}</li>"
        html_content += "</ul>"
        
        if d.get("diseasesToRuleOut"):
            html_content += """
                <h4>Diagnósticos a Descartar (Diagnóstico Diferencial)</h4>
                <ul>
            """
            for rule_out in d["diseasesToRuleOut"]:
                html_content += f"<li>{rule_out}</li>"
            html_content += "</ul>"
            
        if d.get("whenToConsult"):
            html_content += """
                <h4>¿Cuándo Consultar con la Especialista?</h4>
                <ul>
            """
            for w in d["whenToConsult"]:
                html_content += f"<li>{w}</li>"
            html_content += "</ul>"
            
        if d.get("faqs"):
            html_content += "<h4>Preguntas Frecuentes sobre esta Condición</h4>"
            for faq in d["faqs"]:
                html_content += f"""
                <div class="faq-box">
                    <div class="faq-q">¿{faq['question']}?</div>
                    <div class="faq-a">{faq['answer']}</div>
                </div>
                """
                
        if d.get("sources"):
            html_content += """
                <h4>Fuentes de Referencia</h4>
                <ul class="sources-list">
            """
            for src in d["sources"]:
                html_content += f"<li>{src}</li>"
            html_content += "</ul>"
            
        html_content += "</div>"
        
    html_content += """
        </section>

        <!-- SECCIÓN 5: MEDICINA TRANSFUSIONAL -->
        <section id="transfusion">
            <h2>5. Medicina Transfusional y Aféresis Terapéutica</h2>
            <p>Compilación del soporte transfusional y tratamientos avanzados para pacientes de hematología pediátrica.</p>
    """
    
    for t in transfusion:
        html_content += f"""
            <div style="margin-top: 3rem; border-top: 1px solid var(--border); padding-top: 2rem;">
                <h3 style="font-family: 'Playfair Display', serif; text-transform: none; letter-spacing: 0; font-size: 1.5rem; color: var(--primary);">{t['title']}</h3>
        """
        if t.get("intro"):
            html_content += f"<p>{t['intro']}</p>"
            
        html_content += "<h4>Preguntas Frecuentes y Respuestas Clínicas</h4>"
        for sec in t["sections"]:
            html_content += f"""
            <div class="faq-box">
                <div class="faq-q">{sec['q']}</div>
                <div class="faq-a">{sec['a'].replace(chr(10), '<br>')}</div>
            </div>
            """
            
        if t.get("posiblesCausas"):
            html_content += """
                <h4>Causas e Indicaciones Asociadas</h4>
                <ul>
            """
            for pc in t["posiblesCausas"]:
                html_content += f"<li>{pc}</li>"
            html_content += "</ul>"
            
        if t.get("alertQuestion"):
            html_content += f"""
                <div class="callout">
                    <div class="callout-title">{t['alertQuestion']['q']}</div>
                    <div class="callout-text">{t['alertQuestion']['a']}</div>
                </div>
            """
            
        if t.get("sources"):
            html_content += """
                <h4>Referencias Científicas</h4>
                <ul class="sources-list">
            """
            for src in t["sources"]:
                html_content += f"<li>{src}</li>"
            html_content += "</ul>"
            
        html_content += "</div>"
        
    html_content += f"""
        </section>

        <!-- SECCIÓN 6: LEGALES Y CONTACTO -->
        <section id="legales">
            <h2>6. Páginas de Soporte Legal y de Contacto</h2>
            
            <div style="margin-top: 2rem;">
                <h3>{privacy['title']}</h3>
                <p>{privacy['intro']}</p>
    """
    
    for priv_sec in privacy["sections"]:
        html_content += f"""
                <h4>{priv_sec['title']}</h4>
                <p>{priv_sec['content'].replace(chr(10), '<br>')}</p>
        """
        
    html_content += f"""
                <p style="font-size: 0.85rem; color: #777777; margin-top: 2rem;">Última actualización: {privacy['lastUpdate']}</p>
            </div>
            
            <div style="margin-top: 3rem; border-top: 1px solid var(--border); padding-top: 2rem;">
                <h3>Información de Contacto y Agenda</h3>
                <p>Resumen de los canales de atención de la Dra. Lizbeth Hernández:</p>
                <table>
                    <thead>
                        <tr>
                            <th>Concepto</th>
                            <th>Ubicación y Canales de Citas</th>
                        </tr>
                    </thead>
                    <tbody>
                        <tr><td>Doctora</td><td>{doctor['title']} {doctor['name']}</td></tr>
                        <tr><td>Ubicación del Consultorio</td><td>{doctor['address']}, {doctor['city']}, {doctor['state']}</td></tr>
                        <tr><td>WhatsApp Citas</td><td>{doctor['whatsapp']}</td></tr>
                        <tr><td>Teléfono de Contacto</td><td>{doctor['phone']}</td></tr>
                        <tr><td>Email Oficial</td><td>{doctor['email']}</td></tr>
                        <tr><td>Horarios de Atención</td><td>{doctor['schedule']}</td></tr>
                        <tr><td>Costo de Consulta</td><td>${doctor['consultationPrice']} MXN</td></tr>
                        <tr><td>Métodos de Pago Aceptados</td><td>{", ".join(doctor['paymentMethods'])}</td></tr>
                        <tr><td>Enlace Google Maps</td><td><a href="{doctor['googleMapsUrl']}" target="_blank">{doctor['googleMapsUrl']}</a></td></tr>
                    </tbody>
                </table>
            </div>
        </section>
    </main>
    
    <script>
        // Efecto scroll activo para el menú lateral
        const sections = document.querySelectorAll('section');
        const navLinks = document.querySelectorAll('.nav-links a');
        
        window.addEventListener('scroll', () => {{
            let current = '';
            sections.forEach(section => {{
                const sectionTop = section.offsetTop;
                const sectionHeight = section.clientHeight;
                if (pageYOffset >= (sectionTop - 200)) {{
                    current = section.getAttribute('id');
                }}
            }});
            
            navLinks.forEach(link => {{
                link.classList.remove('active');
                if (link.getAttribute('href').slice(1) === current) {{
                    link.classList.add('active');
                }}
            }});
        }});
    </script>
</body>
</html>
"""
    
    with open(html_output, 'w', encoding='utf-8') as f_html:
        f_html.write(html_content)
    print(f"Documento HTML (.html) guardado correctamente en: {html_output}")

    # ----------------------------------------------------
    # 3. GENERACIÓN DEL DOCUMENTO MARKDOWN (.md)
    # ----------------------------------------------------
    # Genera un documento Markdown limpio y perfectamente seccionado.
    md_output = os.path.join(script_dir, "../../Dra_Lizbeth_Verdugo_Contenido.md")
    
    md_content = f"""# Guía de Contenidos Consolidados del Sitio Web
## {doctor['title']} {doctor['name']}
### Especialista en {doctor['specialty']}

*Este documento consolida y organiza de manera perfecta toda la información de textos, consultas, sintomatología y medicina transfusional que compone el sitio web.*

---

## 1. Perfil Profesional de la Doctora

### Biografía
{doctor['bio']}

### Filosofía Médica
> "{doctor['philosophy']}"

### Información General de Consulta
* **Dirección:** {doctor['address']}
* **Teléfono:** {doctor['phone']}
* **WhatsApp:** {doctor['whatsapp']}
* **Email:** {doctor['email']}
* **Precio de Consulta:** ${doctor['consultationPrice']} MXN
* **Métodos de Pago:** {", ".join(doctor['paymentMethods'])}
* **Horarios de Cita:** {doctor['schedule']}
* **Cédula Profesional:** {doctor['cedula']}
* **Cédula de Especialidad:** {doctor['cedulaEspecialidad']}
* **Cédula de Subespecialidad:** {doctor['cedulaSubespecialidad']}
* **Acreditación COFEPRIS:** {doctor['cofepris']}

### Educación
"""
    for edu in doctor["education"]:
        year_str = f" ({edu['year']})" if edu.get("year") else ""
        md_content += f"* **{edu['degree']}** — {edu['institution']}{year_str}\n"
        
    md_content += "\n### Certificaciones\n"
    for cert in doctor["certifications"]:
        year_str = f" ({cert['year']})" if cert.get("year") else ""
        md_content += f"* **{cert['name']}** — {cert['institution']}{year_str}\n"
        
    md_content += "\n### Experiencia Laboral\n"
    for exp in doctor["experience"]:
        md_content += f"* **{exp['year']} - {exp['title']}:** {exp['description']}\n"
        
    md_content += "\n---\n\n## 2. Portafolio de Servicios y Procedimientos Médicos\n"
    
    for idx, s in enumerate(services):
        md_content += f"""
### Servicio {idx+1}: {s['name']}

* **Clasificación:** {s['type'].upper()}
* **Descripción:** {s['description']}
* **Detalle del Procedimiento:** {s['longDescription']}
"""
        if s.get("painDescription"):
            md_content += f"* **Control de Dolor:** {s['painDescription']}\n"
            
        md_content += f"""
#### Ficha Técnica
* **Tipo:** {s['type'].upper()}
* **Anestesia:** {s['anesthesiaType']}
* **Duración:** {s['duration']}
* **Recuperación:** {s['recoveryTime']}
* **¿Es doloroso?:** {"Sí" if s['isPainful'] else "No"}
* **Costo/Precio:** {s['priceRange']}
"""
        for k, v in s["technicalSpecs"].items():
            md_content += f"* **{k}:** {v}\n"
            
        md_content += f"\n#### {s['benefitsTitle']}\n"
        if s.get("benefitsIntro"):
            md_content += f"{s['benefitsIntro']}\n"
        for b in s["benefits"]:
            md_content += f"* {b}\n"
            
        md_content += f"\n#### {s['recommendationsTitle']}\n"
        for rec in s["postOpRecommendations"]:
            md_content += f"* {rec}\n"
            
        if s.get("risks"):
            md_content += f"\n#### {s['risksTitle']}\n"
            for r in s["risks"]:
                md_content += f"* {r}\n"
                
        if s.get("additionalQuestions"):
            md_content += "\n#### Preguntas Frecuentes del Servicio\n"
            for q_item in s["additionalQuestions"]:
                md_content += f"* **¿{q_item['question']}?**\n  *Respuesta:* {q_item['answer']}\n"
                
        if s.get("ctaQuestion"):
            md_content += f"\n> **{s['ctaQuestion']}**\n> *{s['ctaAnswer']}*\n"
            
        md_content += "\n---\n"
        
    md_content += "\n## 3. Guía de Síntomas y Motivos de Consulta\n"
    for idx, sym in enumerate(symptoms):
        md_content += f"""
### Síntoma {idx+1}: {sym['name']}
{sym['description']}

#### {sym['manifestacionesTitle']}
"""
        for man in sym["manifestaciones"]:
            md_content += f"* {man}\n"
            
        md_content += f"\n#### {sym['causesTitle']}\n"
        if sym.get("causesIntro"):
            md_content += f"{sym['causesIntro']}\n"
        for cs in sym["causes"]:
            md_content += f"* {cs}\n"
            
        md_content += f"\n#### ¿Cuándo se debe consultar?\n{sym['whyConsult']}\n"
        
        if sym.get("signosAlarma"):
            md_content += f"\n> **¡ATENCIÓN! {sym['signosAlarmaTitle'] if sym.get('signosAlarmaTitle') else 'Signos de Alarma Hematológicos'}**\n"
            for alarm in sym["signosAlarma"]:
                md_content += f"> * {alarm}\n"
                
        if sym.get("preguntaPrincipal"):
            md_content += f"\n> **{sym['preguntaPrincipal']}**\n> *{sym.get('preguntaPrincipalResponse', 'Consulte para una valoración integral.')}*\n"
            
        md_content += "\n---\n"

    md_content += "\n## 4. Enfermedades y Condiciones de la Sangre\n"
    for idx, d in enumerate(diseases):
        tech_str = f" (*{d['technicalName']}*)" if d.get("technicalName") else ""
        md_content += f"""
### Condición {idx+1}: {d['name']}{tech_str}
**Categoría:** {d['category'].replace('-', ' ').title()}

{d['description']}

#### Síntomas
"""
        for sym_d in d["symptoms"]:
            md_content += f"* {sym_d}\n"
            
        md_content += "\n#### Causas\n"
        for cs in d["causes"]:
            md_content += f"* {cs}\n"
            
        if d.get("riskFactors"):
            md_content += "\n#### Factores de Riesgo\n"
            for risk in d["riskFactors"]:
                md_content += f"* {risk}\n"
                
        if d.get("mexicoStats"):
            md_content += f"\n#### Estadísticas en México\n{d['mexicoStats']}\n"
            
        if d.get("complications"):
            md_content += "\n#### Complicaciones\n"
            for comp in d["complications"]:
                md_content += f"* {comp}\n"
                
        md_content += "\n#### Tratamientos de Soporte y Control\n"
        for tr in d["treatments"]:
            md_content += f"* {tr}\n"
            
        if d.get("diseasesToRuleOut"):
            md_content += "\n#### Diagnósticos Diferenciales a Descartar\n"
            for rule_out in d["diseasesToRuleOut"]:
                md_content += f"* {rule_out}\n"
                
        if d.get("whenToConsult"):
            md_content += "\n#### ¿Cuándo Consultar al Hematólogo Pediatra?\n"
            for w in d["whenToConsult"]:
                md_content += f"* {w}\n"
                
        if d.get("faqs"):
            md_content += "\n#### Preguntas Frecuentes (FAQs)\n"
            for faq in d["faqs"]:
                md_content += f"* **¿{faq['question']}?**\n  *Respuesta:* {faq['answer']}\n"
                
        if d.get("sources"):
            md_content += "\n#### Referencias Médicas\n"
            for src in d["sources"]:
                md_content += f"* {src}\n"
                
        md_content += "\n---\n"

    md_content += "\n## 5. Medicina Transfusional y Aféresis Terapéutica\n"
    for t in transfusion:
        md_content += f"""
### {t['title']}
{t.get('intro', '')}

#### Preguntas y Respuestas
"""
        for sec in t["sections"]:
            md_content += f"* **{sec['q']}**\n  {sec['a']}\n\n"
            
        if t.get("posiblesCausas"):
            md_content += "#### Indicaciones Médicas\n"
            for pc in t["posiblesCausas"]:
                md_content += f"* {pc}\n"
                
        if t.get("alertQuestion"):
            md_content += f"\n> **{t['alertQuestion']['q']}**\n> *{t['alertQuestion']['a']}*\n"
            
        if t.get("sources"):
            md_content += "\n#### Referencias Bibliográficas\n"
            for src in t["sources"]:
                md_content += f"* {src}\n"
                
        md_content += "\n---\n"

    md_content += f"""
## 6. Páginas de Soporte Legal y de Contacto

### {privacy['title']}
{privacy['intro']}
"""
    for priv_sec in privacy["sections"]:
        md_content += f"\n#### {priv_sec['title']}\n{priv_sec['content']}\n"
        
    md_content += f"""
*Última actualización: {privacy['lastUpdate']}*

---

### Página de Contacto e Información General
* **Médico Responsable:** {doctor['title']} {doctor['name']}
* **Ubicación:** {doctor['address']}
* **WhatsApp:** {doctor['whatsapp']}
* **Teléfono Urgencias:** {doctor['phone']}
* **Email:** {doctor['email']}
* **Horario de Citas:** {doctor['schedule']}
* **Inversión por Consulta:** ${doctor['consultationPrice']} MXN
* **Métodos de Pago:** {", ".join(doctor['paymentMethods'])}
* **Enlace Google Maps:** {doctor['googleMapsUrl']}
"""
    
    with open(md_output, 'w', encoding='utf-8') as f_md:
        f_md.write(md_content)
    print(f"Documento Markdown (.md) guardado correctamente en: {md_output}")


if __name__ == '__main__':
    main()
